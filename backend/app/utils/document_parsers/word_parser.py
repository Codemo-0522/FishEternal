"""
Word文档解析器
支持 .docx 和 .doc 格式
"""

import logging
from typing import List, Dict, Any
from .base import DocumentParser, ParseResult

logger = logging.getLogger(__name__)


class WordDocumentParser(DocumentParser):
    """Word文档解析器"""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']
    
    @property
    def parser_name(self) -> str:
        return "word_parser"
    
    def parse_sync(self, content: bytes, filename: str, **kwargs) -> ParseResult:
        """
        同步解析Word文档（基类会自动在线程池中执行）
        
        直接调用阻塞的python-docx等库，无需担心阻塞事件循环
        """
        import os
        file_ext = os.path.splitext(filename)[1].lower()
        
        try:
            if file_ext == '.docx':
                return self._parse_docx(content, filename, **kwargs)
            elif file_ext == '.doc':
                return self._parse_doc(content, filename, **kwargs)
            else:
                return ParseResult.error_result(
                    f"不支持的Word文档格式: {file_ext}",
                    metadata=self._extract_basic_metadata(filename)
                )
                
        except Exception as e:
            logger.error(f"解析Word文档失败: {filename}, 错误: {str(e)}")
            return ParseResult.error_result(
                f"解析Word文档失败: {str(e)}",
                metadata=self._extract_basic_metadata(filename)
            )
    
    def _parse_docx(self, content: bytes, filename: str, **kwargs) -> ParseResult:
        """解析.docx文件（同步方法）"""
        try:
            from docx import Document
            from io import BytesIO
            
            # 从字节流创建Document对象
            doc_stream = BytesIO(content)
            doc = Document(doc_stream)
            
            # 提取文本内容
            text_parts = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # 合并文本
            full_text = "\n".join(text_parts)
            
            # 构建元数据
            metadata = self._extract_basic_metadata(filename)
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "text_length": len(full_text),
                "format": "docx"
            })
            
            # 提取文档属性
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
                if core_props.created:
                    metadata["created"] = core_props.created.isoformat()
                if core_props.modified:
                    metadata["modified"] = core_props.modified.isoformat()
            except Exception as e:
                logger.warning(f"提取文档属性失败: {str(e)}")
            
            return ParseResult.success_result(full_text, metadata)
            
        except ImportError:
            return ParseResult.error_result(
                "缺少python-docx依赖包，无法解析.docx文件",
                metadata=self._extract_basic_metadata(filename)
            )
        except Exception as e:
            logger.error(f"解析.docx文件失败: {str(e)}")
            return ParseResult.error_result(
                f"解析.docx文件失败: {str(e)}",
                metadata=self._extract_basic_metadata(filename)
            )
    
    def _parse_doc(self, content: bytes, filename: str, **kwargs) -> ParseResult:
        """解析.doc文件（旧格式，同步方法）"""
        try:
            # 尝试使用python-docx2txt（如果可用）
            try:
                import docx2txt
                from io import BytesIO
                
                # docx2txt可以处理一些.doc文件
                doc_stream = BytesIO(content)
                text = docx2txt.process(doc_stream)
                
                if text and text.strip():
                    metadata = self._extract_basic_metadata(filename)
                    metadata.update({
                        "text_length": len(text),
                        "format": "doc",
                        "extraction_method": "docx2txt"
                    })
                    return ParseResult.success_result(text.strip(), metadata)
                
            except ImportError:
                pass
            except Exception as e:
                logger.warning(f"使用docx2txt解析.doc文件失败: {str(e)}")
            
            # 尝试使用antiword（如果系统中安装了）
            try:
                import subprocess
                import tempfile
                import os
                
                # 将内容写入临时文件
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
                    tmp_file.write(content)
                    tmp_file_path = tmp_file.name
                
                try:
                    # 使用antiword提取文本（同步阻塞最长30秒，但在线程池中执行）
                    result = subprocess.run(
                        ['antiword', tmp_file_path],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    if result.returncode == 0 and result.stdout.strip():
                        text = result.stdout.strip()
                        metadata = self._extract_basic_metadata(filename)
                        metadata.update({
                            "text_length": len(text),
                            "format": "doc",
                            "extraction_method": "antiword"
                        })
                        return ParseResult.success_result(text, metadata)
                    
                finally:
                    # 清理临时文件
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                        
            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError) as e:
                logger.warning(f"使用antiword解析.doc文件失败: {str(e)}")
            
            # 如果所有方法都失败，返回错误
            return ParseResult.error_result(
                "无法解析.doc文件。请安装docx2txt包或antiword工具，或将文件转换为.docx格式",
                metadata=self._extract_basic_metadata(filename)
            )
            
        except Exception as e:
            logger.error(f"解析.doc文件失败: {str(e)}")
            return ParseResult.error_result(
                f"解析.doc文件失败: {str(e)}",
                metadata=self._extract_basic_metadata(filename)
            )
